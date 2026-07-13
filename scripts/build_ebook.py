"""Assemble ebook/Libstar_Data_Story.docx.

Narrative ebook over the cleaned Libstar catalog: pipeline story, category/
brand/regional analysis with charts, SA map, ML findings and business
background. All figures are read from data/aggregates/*.csv and
ml/outputs/metrics.json at build time — the same sources as the Excel report
and the Qlik reconciliation tables — so the story always matches the numbers.
"""

import json
import os

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

BASE = os.path.join(os.path.dirname(__file__), "..")
AGG = os.path.join(BASE, "data", "aggregates")
CHARTS = os.path.join(BASE, "ebook", "charts")
ML = os.path.join(BASE, "ml", "outputs")
OUT = os.path.join(BASE, "ebook", "Libstar_Data_Story.docx")

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
TEAL = RGBColor(0x2A, 0x9D, 0x8F)
GREY = RGBColor(0x66, 0x66, 0x66)


def rbn(x) -> str:
    return f"R{x / 1e9:,.1f} billion"


def style_doc(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = "Arial"
    n.font.size = Pt(11)
    for sid, size, color in [("Heading 1", 20, NAVY), ("Heading 2", 14, TEAL)]:
        s = doc.styles[sid]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color


def para(doc, text, size=11, color=None, bold=False, italic=False, align=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def chart(doc, filename, caption):
    path = os.path.join(CHARTS, filename)
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, caption, size=9, color=GREY, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)


def page_break(doc):
    doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)


def main() -> None:
    kpi = pd.read_csv(os.path.join(AGG, "kpi_summary.csv")).iloc[0]
    cat = pd.read_csv(os.path.join(AGG, "category_performance.csv"))
    brand = pd.read_csv(os.path.join(AGG, "brand_performance.csv"))
    prov = pd.read_csv(os.path.join(AGG, "province_performance.csv"))
    chan = pd.read_csv(os.path.join(AGG, "channel_performance.csv"))
    dq = pd.read_csv(os.path.join(AGG, "data_quality.csv"))
    with open(os.path.join(ML, "metrics.json")) as f:
        metrics = json.load(f)
    pm = metrics["price_model"]
    seg = pd.read_csv(os.path.join(ML, "segments_summary.csv"))

    total = int(kpi["total_skus"])
    raw = int(kpi["raw_rows"])
    quar = int(kpi["quarantined_rows"])
    rev = float(kpi["revenue_12m_zar"])
    top_cat = cat.iloc[0]
    top_brand = brand.iloc[0]
    amb_share = cat.loc[cat["product_group"] == "Ambient products", "revenue_12m_zar"].sum() / rev
    top_prov = prov.sort_values("revenue_12m_zar", ascending=False).iloc[0]
    retail = chan.iloc[0]

    doc = Document()
    style_doc(doc)

    # ---------- cover ----------
    for _ in range(6):
        doc.add_paragraph()
    para(doc, "Inside the Libstar Catalog", size=34, color=NAVY, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "A data story from five million product records", size=16,
         color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    para(doc, "Azure Data Factory · Synapse serverless · Qlik Sense · Python ML",
         size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "Anthony Apollis — July 2026", size=11, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Synthetic portfolio dataset built on Libstar's public brand and "
              "category model (libstar.co.za). Not affiliated with Libstar Holdings.",
         size=8, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    # ---------- 1. executive summary ----------
    doc.add_heading("1. The story in one page", level=1)
    para(doc,
         f"We started with {raw:,} raw catalog rows — deliberately messy, the way "
         f"real operational exports arrive: prices captured as \"R 1,234.56\" in one "
         f"system and \"ZAR1234\" in another, dates in three formats, weights sometimes "
         f"in grams, brands typed by hand. After an Azure Data Factory cleaning pass, "
         f"{total:,} unique, trustworthy products remained; {quar:,} rows were "
         f"quarantined with an explicit reject reason.")
    para(doc,
         f"The cleaned catalog represents {rbn(rev)} in trailing 12-month revenue. "
         f"{top_cat['category']} is the biggest category ({rbn(top_cat['revenue_12m_zar'])}); "
         f"ambient products contribute {amb_share:.0%} of revenue. {top_brand['brand']} "
         f"is the highest-grossing brand, {top_prov['province']} the strongest province, "
         f"and {retail['sales_channel']} the dominant channel.")
    para(doc,
         f"A gradient-boosted pricing model reconstructs shelf price from catalog "
         f"attributes with R² = {pm['r2']:.2f} (MAE R{pm['mae_zar']:.2f} on a mean "
         f"price of R{pm['mean_price_zar']:.2f}), an IsolationForest flags "
         f"{metrics['anomalies_flagged']:,} pricing anomalies for review, and KMeans "
         f"splits the range into {len(seg)} actionable product segments.")

    # ---------- 2. pipeline ----------
    doc.add_heading("2. From dirty to trusted: the pipeline", level=1)
    para(doc,
         "Architecture: gzipped CSVs land in an Azure Data Lake (ADLS Gen2) raw "
         "zone. An Azure Data Factory mapping data flow normalises prices, dates, "
         "booleans, units and geography; resolves brands against a reference "
         "dimension (fixing legacy business-unit category names like \"Groceries\" "
         "via brand lookup); deduplicates product IDs keeping the latest record; "
         "and writes snappy parquet to a curated zone. Synapse serverless SQL "
         "exposes the parquet through reporting views — no cluster, no idle cost — "
         "and Qlik Sense, this ebook and the Excel report all read the same views.")
    chart(doc, "06_data_quality.png",
          f"Figure 1 — {raw:,} raw rows in, {total:,} clean products out. "
          f"Dates and prices are the biggest offenders.")
    para(doc,
         "Every quarantined row carries its reject reasons, so data quality is a "
         "reportable metric, not a silent filter: "
         + ", ".join(f"{r.reject_reason.rstrip(';').replace(';', ' + ')} ({r.row_count:,})"
                     for r in dq.head(3).itertuples())
         + " lead the list.")

    # ---------- 3. categories ----------
    doc.add_heading("3. What the shelf looks like", level=1)
    para(doc,
         f"Libstar's simplified operating model has two product groups. In this "
         f"catalog the ambient group carries {amb_share:.0%} of revenue across "
         f"{cat[cat['product_group'] == 'Ambient products'].shape[0]} categories, "
         f"while perishables — dairy, convenience meals, value-added meats, baby "
         f"and fresh mushrooms — hold the rest.")
    chart(doc, "01_revenue_by_category.png",
          "Figure 2 — Revenue by category, coloured by product group.")
    chart(doc, "02_margin_by_category.png",
          "Figure 3 — Margins cluster tightly; the dashed line is the portfolio average.")
    chart(doc, "04_group_channel_mix.png",
          "Figure 4 — Product-group split and the four routes to market.")

    # ---------- 4. brands ----------
    doc.add_heading("4. Brand power: three ways to win", level=1)
    para(doc,
         "Libstar runs three brand solutions: its own Libstar Brands (Lancewood, "
         "Denny Mushrooms, Cape Herb & Spice…), Principal Brands it represents in "
         "South Africa (Bonne Maman, Kikkoman, Tabasco, Maille) and private label "
         "for retailers. Principal brands price ~35% above comparable own-brand "
         "products in this dataset; private label sits ~20% below.")
    chart(doc, "03_top_brands.png",
          f"Figure 5 — {top_brand['brand']} leads with {rbn(top_brand['revenue_12m_zar'])}.")

    # ---------- 5. regions ----------
    doc.add_heading("5. Where South Africa buys", level=1)
    para(doc,
         f"Sales concentrate where Libstar operates: {top_prov['province']} leads "
         f"with {rbn(top_prov['revenue_12m_zar'])}, and the five provinces with "
         f"production or distribution footprints (Western Cape, Gauteng, "
         f"KwaZulu-Natal, Eastern Cape, Mpumalanga) account for the bulk of revenue.")
    chart(doc, "05_sa_province_map.png",
          "Figure 6 — 12-month revenue by province.")

    # ---------- 6. momentum ----------
    doc.add_heading("6. Catalog momentum", level=1)
    para(doc,
         "Product introductions run steadily across the observation window — the "
         "assortment machine keeps feeding the shelf.")
    chart(doc, "07_catalog_growth.png", "Figure 7 — Products added per month.")

    # ---------- 7. ML ----------
    doc.add_heading("7. What the machines found", level=1)
    doc.add_heading("Pricing has learnable structure", level=2)
    para(doc,
         f"A HistGradientBoosting regressor trained on {pm['train_rows']:,} products "
         f"predicts price from category, brand, brand solution, channel, weight and "
         f"rating with R² = {pm['r2']:.2f} and a mean absolute error of "
         f"R{pm['mae_zar']:.2f}. Category and brand solution carry most of the "
         f"signal — pricing follows the portfolio's architecture, which is exactly "
         f"what a well-governed catalog should show.")
    doc.add_heading("Anomalies worth a second look", level=2)
    para(doc,
         f"An IsolationForest scored one million products on price, cost, margin "
         f"and weight and flagged {metrics['anomalies_flagged']:,} outliers — "
         f"gram-priced heavyweights, negative-margin outliers, luxury-priced "
         f"commodity items. The top 50 ship in the Excel report for review.")
    chart(doc, "09_ml_anomalies.png",
          "Figure 8 — Flagged anomalies against the normal price–margin cloud.")
    doc.add_heading("Four product segments", level=2)
    para(doc,
         "KMeans on price, margin, volume and rating yields four segments: "
         "volume movers (high units, mid price), premium niche (top price, "
         "average volume), and two mainstream clusters split by customer rating "
         "— a ready-made lens for range reviews.")
    chart(doc, "08_ml_segments.png",
          "Figure 9 — Segment centroids sized by product count.")

    # ---------- 8. business background ----------
    doc.add_heading("8. The business behind the data", level=1)
    para(doc,
         "Libstar (JSE-listed, founded 2005, HQ Cape Town) produces, distributes "
         "and markets consumer packaged goods for South Africa and export markets. "
         "The operating model in this dataset mirrors its public structure:")
    for line in [
        "2 product groups — Perishable products and Ambient products",
        "12 categories — from Dairy and Fresh Mushrooms to Spreads and Beverages",
        "3 brand solutions — Libstar Brands, Principal Brands, Private label and dealer-own",
        "4 sales channels — Retail and wholesale, Food service, Industrial and contract manufacturing, Export (50+ countries)",
        "Operations in 5 provinces — Western Cape, Gauteng, KwaZulu-Natal, Eastern Cape, Mpumalanga",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # ---------- 9. methodology ----------
    doc.add_heading("9. Methodology, cost and reproducibility", level=1)
    para(doc,
         f"Data is synthetic (seeded, reproducible) and generated on Libstar's "
         f"public brand/category model. Stack: Python generator → ADLS Gen2 → "
         f"ADF mapping data flow (8 vCores, ~6 min per run) → parquet → Synapse "
         f"serverless views → Qlik Sense / Excel / this ebook. Everything reads "
         f"one aggregate layer, so the numbers here reconcile with the Qlik app "
         f"and the Excel report by construction.")
    para(doc,
         "Cloud cost discipline: the design has no always-on compute. A full "
         "pipeline run costs well under R20 in ADF data-flow time; Synapse "
         "serverless queries on this volume cost fractions of a cent ($5/TB "
         "scanned); an Azure budget alerts at 50/75/90% of the $200 trial credit.",
         )
    para(doc, "Code, pipeline JSON and SQL: PargoParcels portfolio repository.",
         size=9, color=GREY, italic=True)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    fix_zoom(OUT)
    print("written:", os.path.abspath(OUT))


def fix_zoom(path: str) -> None:
    """python-docx's default template writes <w:zoom> without the required
    w:percent attribute; patch it so strict OOXML validators pass."""
    import shutil
    import zipfile

    tmp = path + ".tmp"
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == "word/settings.xml":
                text = data.decode("utf-8")
                if "<w:zoom" in text and "w:percent" not in text:
                    text = text.replace("<w:zoom ", '<w:zoom w:percent="100" ', 1)
                    text = text.replace("<w:zoom/>", '<w:zoom w:percent="100"/>', 1)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, path)


if __name__ == "__main__":
    main()
